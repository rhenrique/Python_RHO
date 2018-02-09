import pyautogui
import time
import sys
from ctypes import *
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import smtplib
import mimetypes
from email.mime.multipart import MIMEMultipart
from email import encoders
from email.message import Message
from email.mime.audio import MIMEAudio
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.text import MIMEText
import shutil
import subprocess
from datetime import timedelta
from datetime import date
import os


n = []
l = []
a = []
g = []
global login
global w

num = 0

fdate = time.strftime("%Y%m%d")
fl = "CreateUser_LIM" + fdate + ".docx" 
cdate = date.today() + timedelta(days=5)

document = Document()
f = open(fl, 'w')

document.add_heading('Maxion Wheels - Create user in Linux', 0)
document.add_paragraph('To-Do List', style='Heading 1').bold = True

p = document.add_paragraph(' - These users should have ')
p.add_run('gid=60005(qad)').bold = True
p.add_run(' and ')
p.add_run('groups=60005(qad).\n').bold = True
p.add_run(' - Copy the')
p.add_run(' .bash_profile ').bold = True
p.add_run(' and ')
p.add_run('.profile').bold = True
p.add_run(' files from')
p.add_run(' /usr/br/nascimra ').bold = True
p.add_run(' to each new user created, or our start script will not work.\n')

table = document.add_table(1, 5, style='Light List Accent 1')

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'User (Login)'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Site'
hdr_cells[3].text = 'Home Directory'
hdr_cells[4].text = 'Servers'

def quest(num):
	nome = raw_input("Digite o nome da pessoa:\n ")
	n.append(nome)
	login = raw_input("Digite o login da pessoa:\n ")
	l.append(login)
	
	recordset =	[
		{
			"user" : login,
			"name": "(LIM) " + nome,
			"site": "LIM",
			"home": "/users/br",
			"servers": "mxw-qadpl01 / mxw-qaddl01"
		}
	]


	for item in recordset:
		row_cells = table.add_row().cells
		row_cells[0].text = item["user"]
		row_cells[1].text = item["name"]
		row_cells[2].text = item["site"]
		row_cells[3].text = item["home"]
		row_cells[4].text = item["servers"]
		
	amb = raw_input("Escolha um ambiente:\n - [128] - ACO;\n - [130] - Aluminio.\n")
	while amb != '128' and amb != '130':
		amb = raw_input("Escolha entre [128] e [130]:\n")
	a.append(amb)
	askgrupo = raw_input("Deseja adicionar grupo para a pessoa?\n- [s] - SIM;\n- [n] - NAO\n ")
	while askgrupo != 's' and askgrupo != 'n':
		askgrupo = raw_input('Digite [s] ou [n]:\n')
	if askgrupo == 's':
		login = str(l[num]) + '.txt'
		w = open(login, 'w')
		while askgrupo == 's':
			grupo = raw_input("Digite o grupo da pessoa:\n ")
			w.write(grupo + '\n')
			askgrupo = raw_input("Deseja continuar adicionando?\n- [s] - SIM;\n- [n] - NAO\n ")
			while askgrupo != 's' and askgrupo != 'n':
				askgrupo = raw_input('Digite [s] ou [n]:\n')
			if askgrupo == 'n':
				w.close()
				small_file = open(login,'r')
				long_file = open('grupos.txt','r')
				output_file = open(str(l[num]) + '_QAD.txt','w')

				try:
					small_lines = small_file.readlines()
					small_lines_cleaned = [line.rstrip().lower() for line in small_lines]
					long_file_lines = long_file.readlines()
					long_lines_cleaned = [line.rstrip().lower() for line in long_file_lines]
					
					for line in small_lines_cleaned:
						if line in long_lines_cleaned:
							output_file.writelines(line + '\n')
							
				finally:
					small_file.close()
					long_file.close()
					output_file.close()

def ask():
    newline = raw_input("Adicionar mais usuarios?\n- [s] - para SIM;\n- [n] - para NAO.\n")
    return newline

quest(num)

		
yn = ask()
while (yn == 's'):
	num = num + 1
	quest(num)
	yn = ask()
	if (yn == 'n'):
		break
	else:
		True

time.sleep(1)
pyautogui.hotkey('winleft','d')

#Abrindo o QAD
pyautogui.press(['Q','A','D','-','P','R','O','D'])
pyautogui.press('enter')

#Logando no QAD
time.sleep(10)
pyautogui.typewrite('oliveirh')
pyautogui.press('enter')
time.sleep(1)
pyautogui.typewrite('312414ra@#')  
pyautogui.press('enter')
time.sleep(3)
pyautogui.press('1')
pyautogui.press('enter')

if a[num] == '128':
	pyautogui.press('1')
	pyautogui.press('enter')
else:
	pyautogui.press('2')
	pyautogui.press('enter')

num = 0
for item in n:
	time.sleep(3)
	if num > 0:
		pyautogui.typewrite(a[num])
		pyautogui.press('enter')
		time.sleep(3)
	pyautogui.press('enter')
	time.sleep(3)
	if num == 0:
		pyautogui.press('enter')
		time.sleep(3)
	pyautogui.typewrite('36.3.1')
	pyautogui.press('enter')
	time.sleep(3)
	pyautogui.typewrite(l[num])
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.typewrite(n[num])
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.typewrite('po')
	pyautogui.press(['enter','enter'])
	time.sleep(1)
	pyautogui.typewrite('brs')
	pyautogui.press(['enter','enter','enter'])
	time.sleep(1)
	pyautogui.typewrite('PRIMARY')
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.typewrite('GMT-2')
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.typewrite('mailx')
	pyautogui.press(['enter','enter','enter'])
	time.sleep(1)
	pyautogui.typewrite('s')
	pyautogui.press(['enter','enter','enter'])
	time.sleep(1)
	pyautogui.typewrite('QAD_DEF')
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.press(['enter','enter','enter'])
	time.sleep(1)
	pyautogui.press('enter')
	time.sleep(5)
	pyautogui.typewrite(a[num])
	time.sleep(5)
	pyautogui.press('enter')
	pyautogui.press(['enter','enter'])
	
	# Verifica se o arquivo login_QAD.txt existe
	file = os.path.isfile(l[num]+ '_QAD.txt')
	type(file)
	if file == True:
		print 'entrou aqui?'
		with open(l[num] + '_QAD.txt') as f:
			groupList = f.readlines()
			for line in groupList:
				group = line.rstrip('\n')
				time.sleep(2)
				pyautogui.typewrite(group)
				time.sleep(1)
				pyautogui.press('enter')
				time.sleep(1)
			f.close()
	pyautogui.press('f4')
	time.sleep(1)
	pyautogui.press('f4')
	time.sleep(1)
	pyautogui.typewrite('s')
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.press('f4')
	time.sleep(1)
	pyautogui.press('enter')
	time.sleep(1)
	pyautogui.hotkey('ctrl', 'c')
	time.sleep(1)
	num = num + 1
	
num = 0


document.add_page_break()
document.save(fl)

emailto = "rafael.oliveira@maxionwheels.com"
emailfrom = "QAD@maxionwheels.com"
username = "oliveirh@maxionwheels.com"
password = "312414ra@#"

msg = MIMEMultipart()
msg["From"] = emailfrom
msg["To"] = emailto
msg["Subject"] = "MW - Linux team - Create user in Linux"

body = u"""<html><head> <font face="Calibri" size="3" color = "black"></head> <body> Hello S-24, <br><br>Please, see attached to create user in Linux.<br> Complete it until """ + str(cdate) + """</body></html>"""
#msg.attach(MIMEText(body))
msg.attach(MIMEText(body, 'html', 'utf-8'))

ctype, encoding = mimetypes.guess_type(fl)
if ctype is None or encoding is not None:
    ctype = "application/octet-stream"

maintype, subtype = ctype.split("/", 1)

if maintype == "text":
    fp = open(fl)
    # Note: we should handle calculating the charset
    attachment = MIMEText(fp.read(), _subtype=subtype)
    fp.close()
elif maintype == "image":
    fp = open(fl, "rb")
    attachment = MIMEImage(fp.read(), _subtype=subtype)
    fp.close()
elif maintype == "audio":
    fp = open(fl, "rb")
    attachment = MIMEAudio(fp.read(), _subtype=subtype)
    fp.close()
else:
    fp = open(fl, "rb")
    attachment = MIMEBase(maintype, subtype)
    attachment.set_payload(fp.read())
    fp.close()
    encoders.encode_base64(attachment)
attachment.add_header("Content-Disposition", "attachment", filename=fl)
msg.attach(attachment)

try:
    server = smtplib.SMTP('mail.maxionwheels.com',25)
    server.starttls()
    server.sendmail(emailfrom, emailto, msg.as_string())
    server.quit()
    print "Successfully sent email!"
except Exception:
    print "Error: unable to send email"
    server.quit()

user = "MAX\oliveirh"
pwd = "312414ra@#"
networkPath = r'O:\Controle de Acesso - B5_15\Solicitacao de Acesso - QAD\2017\\'

shutil.copy2(fl, networkPath + fl)
print "The file " + fl +  " created was moved to " + networkPath
time.sleep(5)