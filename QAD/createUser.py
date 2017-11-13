# encoding: utf-8
# encoding: iso-8859-1
# encoding: win-1252
import sys
import os
import time
from ctypes import *
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import smtplib
import mimetypes
from email.mime.multipart import MIMEMultipart
from email import encoders
from email.message import Message
from email.mime.audio import MIMEAudio
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.text import MIMEText
import shutil
import subprocess

fdate = time.strftime("%Y%m%d")
#fl = argv[1]
#fl = raw_input("Por favor, digite o nome do arquivo, conforme o exemplo.\nExample: CreateUser_LIM20170914.docx\n")
fl = "CreateUser_LIM" + fdate + ".docx" 
cdate = raw_input("Digite uma data para quando voce ira precisar que seja completado seu ticket.\nExemplo: mm/dd\n")

document = Document()
f = open(fl, 'w')

document.add_heading('Maxion Wheels - Create user in Linux', 0)

print u"Agora vamos criar o arquivo.... Para isso, vou precisar de algumas informações..."
time.sleep(3)
os.system('cls')
def quest():
	quest01 = raw_input("Digite o nome do usuario: \n")
	os.system('cls')
	quest02 = raw_input("Digite o login do usuario \n")
	os.system('cls')
	emailto = raw_input("Digite o e-mail da MAXION para enviar o documento. \nExemplo: rafael.oliveira@maxionwheels.com\n")
	os.system('cls')
	print u"Por favor, confira as informações digitadas:\n- Nome do Usuário: " + quest01 + u"\n- Login do Usuário: " + quest02 + u"\n- Endereço de E-mail: " + emailto
	time.sleep(5)
	recordset =	[
		{
			"user" : quest02,
			"name": "(LIM) " + quest01,
			"site": "LIM",
			"home": "/users/br",
			"servers": "mxw-qadpl01 / mxw-qaddl01"
		}
	]
	for item in recordset:
		row_cells = table.add_row().cells
		row_cells[0].text = item["user"]
    	row_cells[1].text = item["name"]
    	row_cells[2].text = item["site"]
    	row_cells[3].text = item["home"]
    	row_cells[4].text = item["servers"]
	return emailto

def ask():
    newline = raw_input("Adicionar mais usuarios?\n- [s] - para SIM;\n- [n] - para NAO.\n")
    return newline

document.add_paragraph('To-Do List', style='Heading 1').bold = True
p = document.add_paragraph(' - These users should have ')
p.add_run('gid=60005(qad)').bold = True
p.add_run(' and ')
p.add_run('groups=60005(qad).\n').bold = True
p.add_run(' - Copy the')
p.add_run(' .bash_profile ').bold = True
p.add_run(' and ')
p.add_run('.profile').bold = True
p.add_run(' files from')
p.add_run(' /usr/br/nascimra ').bold = True
p.add_run(' to each new user created, or our start script will not work.\n')

table = document.add_table(1, 5, style='Light List Accent 1')

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'User (Login)'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Site'
hdr_cells[3].text = 'Home Directory'
hdr_cells[4].text = 'Servers'

emailto = quest()
yn = ask()
while (yn == 's'):
      quest()
      yn = ask()
      if (yn == 'n'):
          break
      else:
          True

document.add_page_break()
document.save(fl)

emailfrom = "QAD@maxionwheels.com"
username = "oliveirh@maxionwheels.com"
password = "312414ra@#"

msg = MIMEMultipart()
msg["From"] = emailfrom
msg["To"] = emailto
msg["Subject"] = "MW - Linux team - Create user in Linux"

body = u"""<html><head> <font face="Calibri" size="3" color = "black"></head> <body> Hello S-24, <br><br>Please, see attached to create user in Linux.<br> Complete it until """ + cdate + """</body></html>"""
#msg.attach(MIMEText(body))
msg.attach(MIMEText(body, 'html', 'utf-8'))

ctype, encoding = mimetypes.guess_type(fl)
if ctype is None or encoding is not None:
    ctype = "application/octet-stream"

maintype, subtype = ctype.split("/", 1)

if maintype == "text":
    fp = open(fl)
    # Note: we should handle calculating the charset
    attachment = MIMEText(fp.read(), _subtype=subtype)
    fp.close()
elif maintype == "image":
    fp = open(fl, "rb")
    attachment = MIMEImage(fp.read(), _subtype=subtype)
    fp.close()
elif maintype == "audio":
    fp = open(fl, "rb")
    attachment = MIMEAudio(fp.read(), _subtype=subtype)
    fp.close()
else:
    fp = open(fl, "rb")
    attachment = MIMEBase(maintype, subtype)
    attachment.set_payload(fp.read())
    fp.close()
    encoders.encode_base64(attachment)
attachment.add_header("Content-Disposition", "attachment", filename=fl)
msg.attach(attachment)

try:
    server = smtplib.SMTP('mail.maxionwheels.com',25)
    server.starttls()
    server.sendmail(emailfrom, emailto, msg.as_string())
    server.quit()
    print "Successfully sent email!"
except Exception:
    print "Error: unable to send email"
    server.quit()

user = "MAX\oliveirh"
pwd = "312414ra@#"
networkPath = r'O:\Controle de Acesso - B5_15\Solicitacao de Acesso - QAD\2017\\'

shutil.copy2(fl, networkPath + fl)
print "The file " + fl +  " created was moved to " + networkPath
time.sleep(5)
